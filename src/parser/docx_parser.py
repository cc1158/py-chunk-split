import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Dict, Any, Tuple
from .base import BaseParser, Section
from .structure_detector import StructureDetector, create_detector
import re


class DOCXParser(BaseParser):
    """
    DOCX文件解析器

    优先使用原生样式（Heading 1/2/3等），无样式时使用结构特征检测
    支持多语言，不依赖具体格式
    表格和图片会作为完整结构保留
    """

    # Heading样式关键词（多语言）
    HEADING_KEYWORDS = ['heading', 'title', 'caption', 'toc']

    def get_supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']

    def parse(self, file_path: str) -> List[Section]:
        """解析DOCX文件并提取章节"""
        sections = []
        doc = docx.Document(file_path)

        current_section = None
        current_content = []

        # 遍历文档body中的元素（段落和表格交替）
        for elem in doc.element.body:
            if isinstance(elem, CT_P):
                # 这是一个段落
                para = Paragraph(elem, doc)
                style_name = para.style.name if para.style else ''

                # 检查段落中是否包含图片
                image_count = self._has_images_in_paragraph(elem)
                has_images = image_count > 0

                if not has_images:
                    # 没有图片，直接处理整个段落文本
                    text = para.text.strip()
                    if not text:
                        continue

                    # 按换行符分割处理
                    lines = text.split('\n')
                    first_line = lines[0].strip()
                    remaining_lines = '\n'.join(lines[1:]) if len(lines) > 1 else ''

                    is_heading, heading_level = self._detect_from_style(style_name, first_line)

                    if is_heading:
                        # 保存之前的section
                        if current_section:
                            if current_section.content or current_content or current_section.metadata.get('type') == 'table' or current_section.title != 'Preamble':
                                # 只有当current_content有内容时才用它覆盖；否则保持section原有的content
                                if current_content:
                                    current_section.content = '\n'.join(current_content).strip()
                                sections.append(current_section)
                        # 创建新section，title是first_line，content是剩余部分
                        current_section = Section(
                            title=first_line,
                            content=remaining_lines,
                            level=heading_level,
                            metadata={
                                'type': self._determine_section_type(first_line, heading_level),
                                'style': style_name,
                                'is_styled_heading': True,
                            }
                        )
                        current_content = []
                    else:
                        # 检查是否是潜在标题（结构检测）
                        detector = create_detector()
                        is_likely_heading, reason = detector.is_likely_heading(first_line)

                        if is_likely_heading and reason not in ["numbered_list_item", "list_number", "looks_like_list_number"]:
                            # 检查是否是子标题（前一section没有content，且当前是编号子标题）
                            is_subheading = (
                                current_section and
                                not current_section.content and
                                not current_content and
                                reason == "numbering_pattern" and
                                re.match(r'^[0-9]+[.、．、)）]', first_line)
                            )

                            if is_subheading:
                                # 作为前一section的content，而不是新section
                                content_to_add = first_line
                                if remaining_lines:
                                    content_to_add = first_line + '\n' + remaining_lines
                                current_content.append(content_to_add)
                            else:
                                # 保存之前的section
                                if current_section:
                                    if current_section.content or current_content or current_section.metadata.get('type') == 'table' or current_section.title != 'Preamble':
                                        # 只有当current_content有内容时才用它覆盖；否则保持section原有的content
                                        if current_content:
                                            current_section.content = '\n'.join(current_content).strip()
                                        sections.append(current_section)
                                level = detector.detect_heading_level(first_line)
                                current_section = Section(
                                    title=first_line,
                                    content=remaining_lines,
                                    level=level,
                                    metadata={
                                        'type': self._determine_section_type(first_line, level),
                                        'style': style_name,
                                        'is_styled_heading': False,
                                        'detection_reason': reason,
                                    }
                                )
                                current_content = []
                        else:
                            # 非标题，整个段落作为content
                            if current_section is None:
                                current_section = Section(
                                    title='',
                                    content='',
                                    level=0,
                                    metadata={'type': 'preamble'}
                                )
                            # 整个文本（包括剩余行）作为content
                            full_text = remaining_lines if remaining_lines else first_line
                            if remaining_lines and first_line:
                                full_text = first_line + '\n' + remaining_lines
                            elif first_line:
                                full_text = first_line
                            current_content.append(full_text)
                else:
                    # 段落包含图片，需要拆分处理
                    # 先保存当前section（包括current_content中的文本）
                    if current_section:
                        if current_section.content or current_content or current_section.metadata.get('type') == 'table' or current_section.title != 'Preamble':
                            # 只有当current_content有内容时才用它覆盖；否则保持section原有的content
                            if current_content:
                                current_section.content = '\n'.join(current_content).strip()
                            sections.append(current_section)
                        current_section = None
                        current_content = []

                    # 遍历 runs，拆分处理
                    runs = self._extract_runs(elem)
                    pending_text = []  # 累积图片前的文本

                    for run_text, run_has_image in runs:
                        if run_has_image:
                            # 保存图片前的文本（作为独立section）
                            if pending_text:
                                text = ' '.join(pending_text)
                                current_section = Section(
                                    title='',
                                    content=text,
                                    level=0,
                                    metadata={'type': 'preamble'}
                                )
                                sections.append(current_section)
                                current_section = None
                                pending_text = []

                            # 图片：创建单独的 image section
                            current_section = Section(
                                title='[图片]',
                                content='[图片]',
                                level=0,
                                metadata={'type': 'image'}
                            )
                            sections.append(current_section)
                            current_section = None
                        else:
                            # 文本：累积起来
                            if run_text.strip():
                                pending_text.append(run_text.strip())

                    # 处理完后，如果有剩余文本，创建section
                    if pending_text:
                        current_section = Section(
                            title='',
                            content=' '.join(pending_text),
                            level=0,
                            metadata={'type': 'preamble'}
                        )

            elif isinstance(elem, CT_Tbl):
                # 这是一个表格 - 表格不能拆分，需要完整保留
                table = Table(elem, doc)

                # 先保存当前section
                pending_heading_title = None
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    # 如果当前section的标题非空但内容为空，说明这是空的章节标题（如"附件7"）
                    # 应该与下面的表格合并，而不是作为独立section
                    if not current_section.content and current_section.title:
                        pending_heading_title = current_section.title
                        # 不保存这个空heading section
                    elif current_section.content or current_section.metadata.get('type') == 'table' or current_section.title != 'Preamble':
                        sections.append(current_section)
                    current_section = None
                    current_content = []

                # 提取表格内容
                table_data = self._extract_table_data(table)
                table_text = self._format_table_as_text(table_data)

                # 将表格作为独立section
                # 如果有pending的heading标题，应该用这个标题（保持与heading的关联）
                section_title = pending_heading_title if pending_heading_title else '[表格]'
                # content中包含title，这样检索时能知道表格属于哪个章节
                if pending_heading_title:
                    section_content = f"{pending_heading_title}\n{table_text}"
                else:
                    section_content = table_text
                current_section = Section(
                    title=section_title,
                    content=section_content,
                    level=0,
                    metadata={
                        'type': 'table',
                        'row_count': len(table_data),
                        'col_count': len(table_data[0]) if table_data else 0,
                    }
                )
                # 保存表格section到sections，然后重置
                sections.append(current_section)
                current_section = None
                current_content = []

        # 添加最后一个section
        if current_section:
            current_section.content = '\n'.join(current_content).strip()
            if current_section.content or current_section.metadata.get('type') == 'table' or current_section.title != 'Preamble':
                sections.append(current_section)

        return sections

    def _extract_table_data(self, table: Table) -> List[List[str]]:
        """从DOCX表格提取数据"""
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows_data.append(row_data)
        return rows_data

    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """
        将表格格式化为标准 Markdown 表格
        """
        if not table_data:
            return ''

        # 计算每列的最大宽度
        col_widths = []
        for row in table_data:
            for i, cell in enumerate(row):
                cell_len = len(cell)
                if i < len(col_widths):
                    col_widths[i] = max(col_widths[i], cell_len)
                else:
                    col_widths.append(cell_len)

        # 格式化每一行
        formatted_rows = []
        for row_idx, row in enumerate(table_data):
            cells = []
            for i, cell in enumerate(row):
                width = col_widths[i] if i < len(col_widths) else len(cell)
                cells.append(cell.ljust(width))

            row_str = '| ' + ' | '.join(cells) + ' |'
            formatted_rows.append(row_str)

            # 在第一行后添加分隔行
            if row_idx == 0:
                separator = '| ' + ' | '.join('-' * w for w in col_widths) + ' |'
                formatted_rows.append(separator)

        return '\n'.join(formatted_rows)

    def _extract_images_from_doc(self, doc) -> List[Dict[str, Any]]:
        """
        从DOCX文档提取图片信息

        Returns:
            [{'type': str, 'placeholder': str}, ...]
        """
        images = []

        # 遍历文档中的所有内联形状（图片）
        try:
            for shape in doc.inline_shapes:
                # 获取图片尺寸
                width = shape.width
                height = shape.height

                # 根据尺寸判断图片类型
                area = float(width) * float(height) if width and height else 0

                # 假设页面宽度约等于 6 inches (参考值)
                page_area_ref = 6 * 6  # 简化计算
                area_ratio = area / page_area_ref if area else 0

                if area_ratio < 0.01:
                    image_type = 'seal'
                elif area_ratio < 0.05:
                    if height and width and float(height) > float(width) * 2:
                        image_type = 'signature'
                    else:
                        image_type = 'seal'
                elif area_ratio > 0.5:
                    image_type = 'photo'
                else:
                    image_type = 'image'

                images.append({
                    'type': image_type,
                    'placeholder': self._format_image_placeholder(image_type),
                })
        except Exception:
            # 忽略无法提取的图片
            pass

        return images

        return images

    def _format_image_placeholder(self, image_type: str) -> str:
        """生成图片占位符文本"""
        placeholders = {
            'signature': '[签名]',
            'seal': '[印章]',
            'photo': '[图片附件]',
            'image': '[图片]'
        }
        return placeholders.get(image_type, '[图片]')

    def _has_images_in_paragraph(self, para_elem) -> int:
        """
        检测段落中包含的图片数量
        通过检查XML中是否存在<w:drawing>元素来判断
        """
        # Check for drawing elements (inline images)
        drawings = para_elem.xpath('.//w:drawing')
        return len(drawings)

    def _extract_runs(self, para_elem) -> List[Tuple[str, bool]]:
        """
        从段落中提取所有 runs，返回 [(文本, 是否包含图片), ...]

        用于将包含图片的段落拆分为多个部分：
        - 图片前的文本
        - 图片本身
        - 图片后的文本
        """
        runs = []
        for r in para_elem.iter():
            if r.tag.endswith('}r'):
                # 这是一个 run
                text_parts = []
                has_image = False

                for child in r:
                    if child.tag.endswith('}t'):
                        # 文本
                        text = child.text or ''
                        text_parts.append(text)
                    elif child.tag.endswith('}drawing'):
                        # 图片
                        has_image = True

                text = ''.join(text_parts)
                if text or has_image:
                    runs.append((text, has_image))

        return runs

    def _detect_from_style(self, style_name: str, text: str) -> Tuple[bool, int]:
        """
        从样式名称判断是否为标题及级别

        Returns:
            (is_heading, level)
        """
        if not style_name:
            return False, 0

        style_lower = style_name.lower()

        # 检查是否是标题样式
        is_heading_style = any(kw in style_lower for kw in self.HEADING_KEYWORDS)

        if not is_heading_style:
            return False, 0

        # 尝试从样式名称提取级别
        # 匹配 "Heading 1", "标题 1", "标题 2" 等
        level_match = None

        # 数字结尾
        import re
        level_match = re.search(r'(\d+)', style_name)

        if level_match:
            level = int(level_match.group(1))
            return True, level

        # 没有数字，默认级别1
        return True, 1

    def _determine_section_type(self, title: str, level: int) -> str:
        """根据标题和级别判断section类型"""
        title_lower = title.lower()

        if 'appendix' in title_lower or '附件' in title:
            return 'appendix'
        if 'chapter' in title_lower or '第' in title and '章' in title:
            return 'chapter'
        if 'article' in title_lower or '第' in title and '条' in title:
            return 'clause'

        if level == 1:
            return 'chapter'
        elif level == 2:
            return 'clause'
        elif title == 'Preamble':
            return 'preamble'
        else:
            return 'section'

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取DOCX文件的元数据"""
        doc = docx.Document(file_path)
        core_props = doc.core_properties

        metadata = {
            'file_name': file_path,
            'file_type': 'docx',
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
        }

        return metadata
